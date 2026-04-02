import shutil

from openpyxl import Workbook
import docx
from docx.shared import Pt
from docx.shared import Inches
import os
import subprocess

from docx import Document
from docx.shared import Cm, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


from openpyxl.drawing.image import Image
#from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Font, Side, Border



class Docs:

    def docs_excell(dat):
        wb = Workbook()
        wb.create_sheet(title="Ma'lumot",index=0)
        sheet = wb["Ma'lumot"]
        sheet['A1']="O'quvchi haqida ma'lumot"
        nom_title = ["id",'Familiyasi', 'Ismi', "Otasining ismi", "Tug'ulgan yili", 'Telefoni 1', 'Telefoni 2', 'Pasport seriyasi va raqami', "Respublika", 'Viloyati', 'Tumani', "MFY", "Ko'chasi uy raqami", 'JSHSHR', 'Nogironlik', "Status", 'Foto','Uz pasport','Zagran pasport',"Murabbiy ID"]

        for i in range(1,len(nom_title)+1):
            cell = sheet.cell(row=i + 1, column=1)
            cell.value = i

        i = 1
        for ind in nom_title:
            cell=sheet.cell(row=i+1, column=2)
            cell.value = ind
            i += 1

        i=1
        for inb in dat:
            cell=sheet.cell(row=i+1, column=3)
            cell.value = inb
            i += 1


        excel_fayl_nomi = dat[1]+"_"+dat[2]+"_"+dat[3]+".xlsx"
        try:
            wb.save(excel_fayl_nomi)
            os.startfile(excel_fayl_nomi)
        except Exception as e:
              xato = "Exel fayl ochiq turibti yoki boshqa xatolik yuz berdi \n"+str(e)
              return xato

    def docs_word(dat, kim):
        # --------------------------------
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)

        # Вставка фото в правом верхнем углу
        photo_paragraph = doc.add_paragraph()
        photo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = photo_paragraph.add_run()

        try:
            if kim == 1:
                run.add_picture(f'murabbiy_photo/{dat[16]}', width=Cm(3), height=Cm(4))
            elif kim == 0:
                run.add_picture(f'fotopath/{dat[16]}', width=Cm(3), height=Cm(4))
        except Exception as e:
            xato = "Fotosurat formati notog'ri! Fotosuratlar jpg, jpeg, png formatida bo'lishi kerak \n" + str(e)
            return xato

        heading = doc.add_heading('МАЪЛУМОТНОМА', level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        heading1 = doc.add_paragraph(f"{dat[1]} {dat[2]} {dat[3]}")
        heading1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(
            '2018 йил 02 октябрдан:\n(Қаерда ишлаши ва лавозими тўлиқ ёзилади)')

        doc.add_paragraph(f'Туғилган йили:\t{dat[4]}')
        doc.add_paragraph(f'Туғилган жойи:\t{dat[9]} вилояти, {dat[10]} тумани')
        doc.add_paragraph('Миллати:\tўзбек')
        doc.add_paragraph('Партиявийлиги:\tйўқ')
        doc.add_paragraph('Маълумоти:\t-')
        doc.add_paragraph('Тамомлаган:\t-')
        doc.add_paragraph('Маълумоти бўйича мутахассислиги:\t-')
        doc.add_paragraph('Илмий даражаси:\tйўқ')
        doc.add_paragraph('Илмий унвони:\tйўқ')
        doc.add_paragraph('Қайси чет тилларини билади:\tрус тили')
        doc.add_paragraph('Ҳарбий (махсус) унвони:\tйўқ')
        doc.add_paragraph('Давлат мукофотлари билан тақдирланганми:\tйўқ')
        doc.add_paragraph('Сайланадиган органлар аъзосими:\tйўқ')

        heading3 = doc.add_heading('МЕҲНАТ ФАОЛИЯТИ', level=2)
        heading3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(
            f'20__-20__ йй. – {dat[9]} вилояти ______________________________________________________________________________')
        doc.add_paragraph(
            f'20__ й. - ҳ.в.  - {dat[9]} вилояти ____________________________________________________________________________')

        doc.add_paragraph('Малумотнома паспорт, диплом ва меҳнат дафтарчасига асосан тўғри ёзилган')
        doc.add_page_break()

        heading4 = doc.add_heading(f'{dat[1]} {dat[2]} {dat[3]}нинг яқин қариндошлари ҳақида МАЪЛУМОТ', level=2)
        heading4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Таблица с родственниками
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Заголовки
        headers = ['Қариндош-лиги', 'Ф.И.Ш.', 'Туғилган йили ва жойи', 'Иш жойи ва лавозими', 'Яшаш жойи']
        hdr_cells = table.rows[0].cells

        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = p.add_run(header)
            run.bold = True

        manzil = f"{dat[9]} вилояти {dat[10]} тумани {dat[11]} МФЙ {dat[12]} кўчаси"
        relatives = [
            ['Отаси', '--', '--', '--', manzil],
            ['Онаси', '--', '--', '--', manzil],
            ['Синглиси', '--', '--', '--', manzil],
            ['Синглиси', '--', '--', '--', manzil],
            ['Турмуш ўртоғи', '--', '--', '--', manzil],
            ['Ўғли', '--', '--', '--', manzil],
            ['Қайнотаси', '--', '--', '--', manzil],
            ['Қайнонаси', '--', '--', '--', manzil]
        ]

        for row_data in relatives:
            row = table.add_row().cells
            for i, item in enumerate(row_data):
                row[i].text = item

        doc.add_paragraph(f'Паспорт: {dat[7]} - {dat[10]} туман ИИБ __.__.__')
        doc.add_paragraph('Диплом: _ № ______    __.__.20__  Рег № 2358')
        doc.add_paragraph('Мех.даф : ______')
        doc.add_paragraph(f'Телефон рақами: {dat[5]}')
        doc.add_paragraph(f'Телефон рақами: {dat[6]}')
        doc.add_page_break()
        # ----------------------------------------------------------
        if kim == 1:
            if dat[17]:
                doc.add_picture(f'murabbiy_pasport_uz/{dat[17]}', width=Inches(6))
            if dat[18]:
                doc.add_picture(f'murabbiy_pasport_zagran/{dat[18]}', width=Inches(6))
            if dat[19]:
                doc.add_picture(f'mukofot1/{dat[19]}', width=Inches(6))
            if dat[20]:
                doc.add_picture(f'mukofot2/{dat[20]}', width=Inches(6))
            if dat[21]:
                doc.add_picture(f'mukofot3/{dat[21]}', width=Inches(6))
        elif kim == 0:
            if dat[17]:
                doc.add_picture(f'pasport_uz/{dat[17]}', width=Inches(6))
            if dat[18]:
                doc.add_picture(f'pasport_zagran/{dat[18]}', width=Inches(6))
            if dat[20]:
                doc.add_picture(f'mukofot1/{dat[20]}', width=Inches(6))
            if dat[21]:
                doc.add_picture(f'mukofot2/{dat[21]}', width=Inches(6))
            if dat[22]:
                doc.add_picture(f'mukofot3/{dat[22]}', width=Inches(6))


        word_fayl_nomi = dat[1] + "_" + dat[2] + "_" + dat[3] + ".docx"
        #os.startfile(word_fayl_nomi)

        try:
            doc.save(word_fayl_nomi)
            os.startfile(word_fayl_nomi)
        except Exception as e:
              xato = "Word fayl ochiq turibti yoki boshqa xatolik yuz berdi \n"+str(e)
              return xato

    def allpupilexcel(dat,kim):
        wb = Workbook()
        wb.create_sheet(title="Barcha o'quvchi sportchilar", index=0)
        sheet = wb["Barcha o'quvchi sportchilar"]
        if kim == 1:
            nom_title = ["id", 'Familiyasi', 'Ismi', "Otasining ismi", "Tug'ulgan yili", 'Telefoni 1', 'Telefoni 2',
                     'Pasport seriyasi va raqami', "Respublika", 'Viloyati', 'Tumani', "MFY", "Ko'chasi uy raqami",
                     'JSHSHR', 'Nogironlik', "Status", 'Foto',"Uz pasport", "Zagran pasport","Mukofot 3","Mukofot 2","Mukofot 3" ]
        elif kim == 0:
            nom_title = ["id", 'Familiyasi', 'Ismi', "Otasining ismi", "Tug'ulgan yili", 'Telefoni 1', 'Telefoni 2',
                         'Pasport seriyasi va raqami', "Respublika", 'Viloyati', 'Tumani', "MFY", "Ko'chasi uy raqami",
                         'JSHSHR', 'Nogironlik', "Status", 'Foto', "Uz pasport", "Zagran pasport", "Murabbiy ID","Mukofot 3",
                         "Mukofot 2", "Mukofot 3"]

        i = 0
        for ind in nom_title:
            cell = sheet.cell(row=1, column=i + 1)
            cell.value = ind
            i += 1

        u = 2
        for ustun in dat:
            i = 0
            for inb in ustun:
                cell = sheet.cell(row=u, column=i+1)
                cell.value = inb
                i += 1
            u +=1

        excel_fayl_nomi = "barchasi.xlsx"
        try:
            wb.save(excel_fayl_nomi)
            os.startfile(excel_fayl_nomi)
        except Exception as e:
            xato = "Exel fayl ochiq turibti yoki boshqa xatolik yuz berdi \n" + str(e)
            return xato


    def todirectory(dat, mo):
        # Путь к новой папке
        #print(dat)
        # Получаем путь к рабочему столу пользователя
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        # Название папки, которую хотим создать
        folder_name = str(dat[0])+"_"+dat[1]+"_"+dat[2]+"_"+dat[3].replace(" ", "_")
        # Полный путь к новой папке
        target_folder = os.path.join(desktop_path, folder_name)
        # Создаём папку, если она НЕ существует. Если существует — выдаём ошибку
        try:
            os.mkdir(target_folder)
            #print(f"Папка '{folder_name}' создана на рабочем столе.")
        except Exception:
            ex= f"'{folder_name}' papkasi oldin yaratilgan ekan!!!"
            return ex
        # Файлы для копирования: (исходный_файл, новое_имя)
        if mo == 0:
            files_to_copy = [
                (f"fotopath/{dat[16]}", f"{dat[0]}_Fotosurat_3x4.jpg"),
                (f"pasport_uz/{dat[17]}", f"{dat[0]}_Pasporti.jpg"),
                (f"pasport_zagran/{dat[18]}", f"{dat[0]}_Zagran_pasporti.jpg"),
                (f"mukofot1/{dat[20]}", f"{dat[0]}_Mukofot1.jpg"),
                (f"mukofot2/{dat[21]}", f"{dat[0]}_Mukofot2.jpg"),
                (f"mukofot3/{dat[22]}", f"{dat[0]}_Mukofot3.jpg")
            ]
        elif mo == 1:
            files_to_copy = [
                (f"murabbiy_photo/{dat[16]}", f"{dat[0]}_Fotosurat_3x4.jpg"),
                (f"murabbiy_pasport_uz/{dat[17]}", f"{dat[0]}_Pasporti.jpg"),
                (f"murabbiy_pasport_zagran/{dat[18]}", f"{dat[0]}_Zagran_pasporti.jpg"),
                (f"mukofot1/{dat[19]}", f"{dat[0]}_Mukofot1.jpg"),
                (f"mukofot2/{dat[20]}", f"{dat[0]}_Mukofot2.jpg"),
                (f"mukofot3/{dat[21]}", f"{dat[0]}_Mukofot3.jpg")
            ]

        # Копирование и переименование
        for source, new_name in files_to_copy:
            if os.path.exists(source):
                destination = os.path.join(target_folder, new_name)
                shutil.copy(source, destination)

        subprocess.Popen(f'explorer "{target_folder}"')

    def to_bayonnoma(self, l1nomi, l2bayon):
        # Создание книги и листа
        wb = Workbook()
        ws = wb.active
        ws.sheet_view.zoomScale = 130  # можно задать 80, 120 и т.д.

        ws.title = "Баённома"
        time_n_r12 = Font(name='Times New Roman', size=12, bold=True)

        # Настройка ширины столбцов для вида таблицы
        ws.column_dimensions['A'].width = 3.14
        ws.column_dimensions['B'].width = 28
        ws.column_dimensions['C'].width = 13.29
        ws.column_dimensions['D'].width = 8.86
        ws.column_dimensions['E'].width = 15.57
        ws.column_dimensions['F'].width = 9.43
        ws.column_dimensions['G'].width = 8.43

        # Вставка изображений
        img_left = Image(
            "test_images/logoparauz1.png")  # Левая картинка (вы прислали только одну, её используем дважды)
        img_right = Image("test_images/laylak.jpg")  # Правая картинка (в реальности должна быть другая)

        # Изменим размер изображения (под ячейки)
        img_left.width = 90
        img_left.height = 100
        img_right.width = 80
        img_right.height = 130

        # Заголовок (объединённые ячейки)
        ws.merge_cells('A1:G1')

        # Вставляем в нужные координаты
        ws.add_image(img_left, "A1")  # Левый верхний угол
        ws.add_image(img_right, "G1")  # Правый верхний угол

        if l1nomi[1] == "1-Klassni tanlang":
            l1nomi[1]=""
        else:
            l1nomi[1] = l1nomi[1]+" "
        if l1nomi[2] == "2-Klassni tanlang":
            l1nomi[2]=""
        else:
            l1nomi[2] = l1nomi[2]+" "
        if l1nomi[3] == "3-Klassni tanlang":
            l1nomi[3]=""
        #print(l1nomi[4])
        if l1nomi[4] == "100м":
            l1nomi[4] = "100 метр масофага\nюгуриш бўйича"
        if l1nomi[4] == "200м":
            l1nomi[4] = "200 метр масофага\nюгуриш бўйича"
        if l1nomi[4] == "400м":
            l1nomi[4] = "400 метр масофага\nюгуриш бўйича"
        if l1nomi[4] == "1500м":
            l1nomi[4] = "1500 метр масофага\nюгуриш бўйича"
        if l1nomi[4] == "Ядро":
            l1nomi[4] = "Ядро улоқтириш\nбўйича"
        if l1nomi[4] == "Диск":
            l1nomi[4] = "Диск улоқтириш\nбўйича"
        if l1nomi[4] == "Клаб":
            l1nomi[4] = "CLUB THROW улоқтириш\nбўйича"
        if l1nomi[4] == "Узунлик":
            l1nomi[4] = "Узунликка сакраш\nбўйича"
        if l1nomi[4] == "Найза":
            l1nomi[4] = "Ядро улоқтириш\nбўйича"
        if l1nomi[4] == "Баландлик":
            l1nomi[4] = "Баландликка сакраш\nбўйича"
        if l1nomi[4] == "Арава":
            l1nomi[4] = "100 метр масофага\nюгуриш бўйича"



        ws['A1'] = f"Пара енгил атлетика бўйича Ўзбекистон чемпионати\n{l1nomi[1]}{l1nomi[2]}{l1nomi[3]} тоифасида {l1nomi[4]} мусобақа\nБАЁННОМАСИ"
        ws['A1'].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        # ws['A1'].font = Font(bold=True)
        ws.row_dimensions[1].height = 90
        ws['A1'].font = time_n_r12
        if l1nomi[0] == "Э":
            jns = "ЭРКАКЛАР"
        else:
            jns = "АЁЛЛАР"
        # Женская категория
        ws.merge_cells('A2:G2')
        ws['A2'] = jns
        ws['A2'].alignment = Alignment(horizontal="center", vertical="center")
        # ws['A2'].font = Font(bold=True)
        ws['A2'].font = time_n_r12
        ws.row_dimensions[2].height = 18

        # Дата и город
        ws.merge_cells('A3:B3')
        ws['A3'] = "6-8 май 2024 йил"
        ws['A3'].alignment = Alignment(horizontal="center", vertical="center")
        ws['A3'].font = time_n_r12

        ws.merge_cells('E3:G3')
        ws['E3'] = "Ангрен шаҳри"
        ws['E3'].alignment = Alignment(horizontal="center", vertical="center")
        ws['E3'].font = time_n_r12

        # ws['B3'].alignment = ws['E3'].alignment = Alignment(horizontal="center", vertical="center")
        # ws['B3'].font = ws['E3'].font = Font(bold=True)

        # Заголовки таблицы
        headers = ['№', 'Ф.И.О', 'Туғилган йили', 'Вилоят', 'Кўкрак рақами', 'Натижа', 'Ўрин']
        thick_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin'))

        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=4, column=col)
            cell.value = header
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            # cell.font = Font(bold=True)
            cell.font = time_n_r12
            cell.border = thick_border

            # Telo таблицы
            thick_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin'))
            # Listni hujjatga moslashtirib olamiz

            bay_uch = [item[:3] + (item[8],) for item in l2bayon]

            # Пишем таблицу, начиная с 6-й строки (row + 5)
            for row_idx, row_data in enumerate(bay_uch, start=1):
                for col_idx, cell_value in enumerate(row_data, start=1):
                    cell = ws.cell(row=row_idx + 4, column=col_idx)
                    cell.value = cell_value
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    cell.font = Font(name='Times New Roman', size=12, bold=False)
                    cell.border = thick_border
                for col_idx, cell_value in enumerate(["","",""], start=1):
                    cell = ws.cell(row=row_idx + 4, column=col_idx+4)
                    cell.value = cell_value
                    cell.border = thick_border

        # Сохраняем файл
        #excel_fayl_nomi = "bayonnoma.xlsx"
        #excel_fayl_nomi = (l1nomi[1]+l1nomi[2]+l1nomi[3]+l1nomi[4][:5]+"_bayonnoma.xlsx").replace(" ",'_')

        # Получаем путь к папке "Документы"
        documents_folder = os.path.join(os.path.expanduser("~"), "Documents")

        # Формируем полный путь к файлу
        excel_fayl_nomi = os.path.join(documents_folder, "bayonnoma.xlsx")

        try:
            wb.save(excel_fayl_nomi)
            os.startfile(excel_fayl_nomi)
        except Exception as e:
           ex = "Bayonnoma fayli excelda ochiq holatda. Yoki boshqa hatolik yuz berdi.\nHatolik: " + str(e)
           return ex























